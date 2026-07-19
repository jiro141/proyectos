#!/usr/bin/env python3
"""
Generador del Manual de Usuario - Pizzería El Patio
Genera un documento .docx profesional con portada, índice, tablas y contenido completo.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Colores corporativos ───
COLOR_PRIMARY = RGBColor(0xF5, 0x73, 0x00)  # Naranja
COLOR_SECONDARY = RGBColor(0x1F, 0x29, 0x37)  # Azul oscuro
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BG_LIGHT = RGBColor(0xF8, 0xF9, 0xFA)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)
COLOR_BORDER = RGBColor(0xDE, 0xE2, 0xE6)

doc = Document()

# ═══════════════════════════════════════════════════════════
# CONFIGURACIÓN GLOBAL
# ═══════════════════════════════════════════════════════════

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = COLOR_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 5):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = 'Calibri'
    h_font.bold = True
    h_font.color.rgb = COLOR_SECONDARY if level > 1 else COLOR_PRIMARY
    if level == 1:
        h_font.size = Pt(22)
        h_style.paragraph_format.space_before = Pt(24)
        h_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h_font.size = Pt(16)
        h_style.paragraph_format.space_before = Pt(20)
        h_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        h_font.size = Pt(13)
        h_style.paragraph_format.space_before = Pt(16)
        h_style.paragraph_format.space_after = Pt(6)
    elif level == 4:
        h_font.size = Pt(11)
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(4)

# ─── Margenes ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


# ═══════════════════════════════════════════════════════════
# FUNCIONES AUXILIARES
# ═══════════════════════════════════════════════════════════

def add_separator():
    """Línea decorativa naranja"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 50)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.size = Pt(8)


def add_info_table(headers, rows):
    """Tabla con formato profesional"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Background naranja
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F57300"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            # Alternar color de fondo
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3EB"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # espacio post-tabla
    return table


def add_bullet(text, bold_prefix=None, level=0):
    """Agrega un elemento de lista con viñeta"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_numbered(text, bold_prefix=None):
    """Agrega un elemento de lista numerada"""
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
    """Párrafo con formato rápido"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_note(text, type="info"):
    """Caja de nota informativa, advertencia o tip"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    labels = {
        "info": "📌 INFORMACIÓN: ",
        "warning": "⚠️ IMPORTANTE: ",
        "tip": "💡 CONSEJO: ",
        "caution": "🔴 PRECAUCIÓN: "
    }
    label = labels.get(type, "📌 ")

    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)

    colors = {
        "info": RGBColor(0x05, 0x6B, 0xE8),
        "warning": RGBColor(0xE8, 0x6B, 0x05),
        "tip": RGBColor(0x05, 0x9E, 0x69),
        "caution": RGBColor(0xDC, 0x35, 0x35)
    }
    run.font.color.rgb = colors.get(type, COLOR_PRIMARY)

    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = COLOR_GRAY

    # Fondo gris claro
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
    p._element.get_or_add_pPr().append(shading)
    return p


def add_bold_normal(bold_text, normal_text):
    """Párrafo con parte en negrita y parte normal"""
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.bold = True
    run_b.font.size = Pt(11)
    run_n = p.add_run(normal_text)
    run_n.font.size = Pt(11)
    return p


def add_code_block(code_text):
    """Bloque de código con fondo gris"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p._element.get_or_add_pPr().append(shading)
    return p


def add_image_if_exists(path, width=5):
    """Agrega imagen si existe el archivo"""
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except:
            return False
    return False


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════

# Espacio superior
for _ in range(6):
    doc.add_paragraph()

# Logo
logo_path = os.path.join(os.path.dirname(__file__), 'frontend', 'public', 'logo.png')
logo_path2 = os.path.join(os.path.dirname(__file__), 'assets', 'logo.png')
if not add_image_if_exists(logo_path, width=3.5):
    add_image_if_exists(logo_path2, width=3.5)

doc.add_paragraph()

# Título principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE USUARIO')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = COLOR_PRIMARY

# Subtítulo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Gestión para Restaurante')
run.font.size = Pt(18)
run.font.color.rgb = COLOR_SECONDARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Pizzería El Patio')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = COLOR_SECONDARY

doc.add_paragraph()
doc.add_paragraph()

# Línea decorativa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('═' * 60)
run.font.color.rgb = COLOR_PRIMARY
run.font.size = Pt(12)

doc.add_paragraph()

# Datos de la portada
portada_data = [
    ('Versión:', '1.0.0'),
    ('Fecha:', datetime.date.today().strftime('%d de %B de %Y')),
    ('Plataforma:', 'Web (React + Django REST) + Android'),
    ('Idioma:', 'Español (Argentina)'),
]

for label, value in portada_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_SECONDARY
    run = p.add_run(value)
    run.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════════
# ÍNDICE / TABLA DE CONTENIDOS
# ═══════════════════════════════════════════════════════════

doc.add_heading('Índice', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(Actualice la tabla de contenidos en Word presionando Ctrl+A → F9 para ver los números de página)')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = COLOR_GRAY

# Insertar campo TOC (Tabla de Contenidos)
p = doc.add_paragraph()
run = p.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._element.append(fldChar1)

run2 = p.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-4" \\h \\z \\u </w:instrText>')
run2._element.append(instrText)

run3 = p.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._element.append(fldChar2)

run4 = p.add_run('[Actualice la tabla de contenidos en Word: Ctrl+A → F9]')
run4.font.color.rgb = COLOR_GRAY
run4.font.size = Pt(10)
run4.italic = True

run5 = p.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._element.append(fldChar3)

page_break()

# ═══════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════

doc.add_heading('1. Introducción', level=1)

add_para(
    'El Sistema de Gestión para Restaurante "Pizzería El Patio" es una aplicación '
    'completa diseñada para administrar todas las operaciones diarias de un restaurante '
    'o pizzería. Proporciona herramientas para la gestión de mesas, pedidos, cocina, '
    'facturación y cierre de caja, todo en tiempo real.'
)

add_para(
    'El sistema está compuesto por un frontend web moderno (React) accesible desde '
    'cualquier navegador o dispositivo móvil, y un backend robusto (Django REST) que '
    'procesa toda la lógica de negocio y almacena los datos de forma segura.'
)

doc.add_heading('1.1 Propósito del Sistema', level=2)
add_bullet('Optimizar la comunicación entre meseros, cocina y caja.')
add_bullet('Agilizar la toma de pedidos y la facturación.')
add_bullet('Mantener un control preciso de ingresos y ventas.')
add_bullet('Generar reportes de ventas y facturas en PDF.')
add_bullet('Notificar en tiempo real los cambios de estado de los pedidos.')

doc.add_heading('1.2 Roles de Usuario', level=2)
add_info_table(
    ['Rol', 'Descripción', 'Acceso Principal'],
    [
        ['Administrador', 'Gestiona el menú, usuarios y configuración del sistema', 'Panel admin, Menú, todo el sistema'],
        ['Mesero', 'Toma pedidos, asigna mesas, entrega pedidos', 'Mesas, Pedidos, Kanban'],
        ['Cocina', 'Prepara los platos y marca estados de preparación', 'Cocina (Kanban de items)'],
        ['Cajero', 'Factura, cobra, genera reportes y cierra caja', 'Caja, Reportes'],
    ]
)

doc.add_heading('1.3 Arquitectura General', level=2)
add_para(
    'El sistema sigue una arquitectura cliente-servidor con comunicación en tiempo real:'
)

add_bold_normal('Frontend (React + Vite): ', 'Aplicación web de página única (SPA) que se ejecuta en el navegador. ')
add_para(
    'Responsable de la interfaz de usuario, la navegación y la experiencia visual. '
    'Se comunica con el backend a través de una API REST y WebSockets.'
)

add_bold_normal('Backend (Django REST Framework): ', 'API REST que procesa todas las solicitudes. ')
add_para(
    'Maneja la lógica de negocio, la autenticación JWT, la base de datos, '
    'las notificaciones push (Firebase) y los WebSockets en tiempo real.'
)

add_bold_normal('Base de Datos (PostgreSQL / SQLite / MySQL): ', 'Almacena toda la información del sistema: ')
add_para(
    'usuarios, mesas, platos del menú, pedidos, facturas y dispositivos registrados.'
)

page_break()

# ═══════════════════════════════════════════════════════════
# 2. REQUISITOS DEL SISTEMA
# ═══════════════════════════════════════════════════════════

doc.add_heading('2. Requisitos del Sistema', level=1)

doc.add_heading('2.1 Para Usuarios (Frontend)', level=2)

add_info_table(
    ['Componente', 'Mínimo', 'Recomendado'],
    [
        ['Navegador Web', 'Chrome 90+, Firefox 90+, Edge 90+', 'Chrome 120+ / Edge 120+'],
        ['Conexión a Internet', '1 Mbps', '5 Mbps o superior'],
        ['Resolución de Pantalla', '1024 × 768', '1366 × 768 o superior'],
        ['Android (App Nativa)', 'Android 8.0 (API 26)', 'Android 12+ (API 31+)'],
        ['JavaScript', 'Habilitado', 'Habilitado'],
    ]
)

doc.add_heading('2.2 Para Despliegue (Backend)', level=2)
add_info_table(
    ['Componente', 'Desarrollo', 'Producción'],
    [
        ['Python', '3.11+', '3.11+'],
        ['Base de Datos', 'SQLite (incluido)', 'PostgreSQL 15 / MySQL 8'],
        ['Servidor', 'Django runserver', 'Gunicorn + Daphne (ASGI)'],
        ['Node.js', '18+ (para build frontend)', '18+ (para build frontend)'],
        ['Docker', 'Opcional', 'Recomendado'],
    ]
)

page_break()

# ═══════════════════════════════════════════════════════════
# 3. INSTALACIÓN Y CONFIGURACIÓN
# ═══════════════════════════════════════════════════════════

doc.add_heading('3. Instalación y Configuración', level=1)

doc.add_heading('3.1 Inicio Rápido con Docker', level=2)
add_para('La forma más rápida de poner el sistema en funcionamiento es usando Docker:')

add_code_block(
    'docker-compose up -d\n'
    '# Frontend:  http://localhost:5173\n'
    '# Backend:   http://localhost:8000\n'
    '# Admin:     http://localhost:8000/admin/'
)

add_note(
    'Asegúrese de tener Docker y Docker Compose instalados en su sistema.',
    'info'
)

doc.add_heading('3.2 Instalación Manual - Backend', level=2)

add_numbered('Clonar el repositorio y navegar al directorio del backend:')
add_code_block('cd backend')

add_numbered('Crear y activar un entorno virtual:')
add_code_block(
    '# Windows:\n'
    'python -m venv venv\n'
    'venv\\Scripts\\activate\n\n'
    '# Linux/Mac:\n'
    'python3 -m venv venv\n'
    'source venv/bin/activate'
)

add_numbered('Instalar las dependencias:')
add_code_block(
    '# Para desarrollo:\n'
    'pip install -r requirements/dev.txt\n\n'
    '# Para producción:\n'
    'pip install -r requirements/prod.txt'
)

add_numbered('Configurar las variables de entorno. Copie el archivo de ejemplo:')
add_code_block(
    '# Windows:\n'
    'copy .env.example .env\n\n'
    '# Linux/Mac:\n'
    'cp .env.example .env'
)

add_para('Edite el archivo .env con sus configuraciones:')
add_code_block(
    'SECRET_KEY=tu-clave-secreta-aqui\n'
    'DEBUG=True\n'
    'ALLOWED_HOSTS=*\n'
    'DB_ENGINE=django.db.backends.sqlite3\n'
    '# Para PostgreSQL:\n'
    '# DB_ENGINE=django.db.backends.postgresql\n'
    '# DB_NAME=pizzeria\n'
    '# DB_USER=pizzeria\n'
    '# DB_PASSWORD=pizzeria\n'
    '# DB_HOST=localhost\n'
    '# DB_PORT=5432'
)

add_numbered('Ejecutar las migraciones de la base de datos:')
add_code_block('python manage.py migrate')

add_numbered('Crear un superusuario (administrador):')
add_code_block('python manage.py createsuperuser')

add_numbered('Iniciar el servidor de desarrollo:')
add_code_block('python manage.py runserver 0.0.0.0:8000')

doc.add_heading('3.3 Instalación Manual - Frontend', level=2)

add_numbered('Navegar al directorio del frontend:')
add_code_block('cd frontend')

add_numbered('Instalar las dependencias de Node.js:')
add_code_block('npm install')

add_numbered('Configurar la URL del backend (opcional, por defecto usa proxy de Vite):')
add_code_block(
    '# Editar .env o .env.production\n'
    'VITE_API_URL=http://localhost:8000'
)

add_numbered('Iniciar el servidor de desarrollo:')
add_code_block('npm run dev')

add_para('La aplicación estará disponible en: http://localhost:5173')

doc.add_heading('3.4 Construir para Producción', level=2)
add_code_block(
    'cd frontend\n'
    'npm run build      # Genera la carpeta dist/\n\n'
    '# El contenido de dist/ se sirve con Nginx, Apache, o Netlify.'
)

page_break()

# ═══════════════════════════════════════════════════════════
# 4. PRIMEROS PASOS
# ═══════════════════════════════════════════════════════════

doc.add_heading('4. Primeros Pasos', level=1)

doc.add_heading('4.1 Acceso al Sistema', level=2)
add_para('Abra su navegador web y diríjase a la URL donde está alojada la aplicación:')
add_bullet('Desarrollo local: http://localhost:5173')
add_bullet('Producción: La URL proporcionada por su administrador')

doc.add_heading('4.2 Pantalla de Inicio de Sesión', level=2)
add_para(
    'Al acceder al sistema, verá la pantalla de inicio de sesión con el logo '
    'de "Pizzería El Patio" y un formulario para ingresar sus credenciales.'
)

add_bold_normal('Credenciales: ', 'Son proporcionadas por el administrador del sistema.')
add_para('')
add_bullet('Usuario: Su nombre de usuario (ej: "juan", "maria_cocina")')
add_bullet('Contraseña: La contraseña asignada por el administrador')

add_para(
    'Presione el botón "Iniciar Sesión" para acceder al sistema. '
    'Si las credenciales son correctas, será redirigido automáticamente '
    'al panel principal según su rol.'
)

add_note(
    'Si olvidó su contraseña, contacte al administrador del sistema para restablecerla.',
    'warning'
)

doc.add_heading('4.3 Panel Principal (Dashboard)', level=2)
add_para(
    'Después de iniciar sesión, verá el panel principal con un mensaje de bienvenida '
    'personalizado que incluye su nombre de usuario y la fecha actual. '
    'En la barra lateral izquierda encontrará las opciones de navegación disponibles '
    'según su rol.'
)

add_bold_normal('Barra Lateral: ', 'Contiene las secciones a las que tiene acceso.')
add_para('')
add_bullet('Icono de usuario: Muestra su nombre y rol actual')
add_bullet('Enlaces de navegación: Varían según el rol (Mesas, Pedidos, Cocina, Caja, Menú)')
add_bullet('Cerrar Sesión: Botón para salir del sistema')

page_break()

# ═══════════════════════════════════════════════════════════
# 5. MANUAL DEL USUARIO - FRONTEND POR ROLES
# ═══════════════════════════════════════════════════════════

doc.add_heading('5. Manual del Usuario', level=1)

add_para(
    'A continuación se detallan las funcionalidades disponibles para cada rol '
    'dentro del sistema. Seleccione la sección correspondiente a su perfil.'
)

# ─── 5.1 MESERO ───
doc.add_heading('5.1 Mesero', level=2)

add_para(
    'El mesero es responsable de tomar pedidos, gestionar las mesas y entregar '
    'los platos a los comensales. Las secciones disponibles son: Mesas y Pedidos.',
    bold=True
)

doc.add_heading('5.1.1 Gestión de Mesas', level=3)
add_para('La página de Mesas muestra un mapa visual de todas las mesas del restaurante. '
         'Cada mesa se representa como una tarjeta con su número y estado actual:')

add_info_table(
    ['Estado', 'Color', 'Descripción'],
    [
        ['Libre', 'Verde', 'Mesa disponible para nuevos comensales'],
        ['Ocupada', 'Rojo / Naranja', 'Mesa con un pedido activo'],
        ['Limpieza', 'Amarillo', 'Mesa siendo limpiada después de pagar'],
    ]
)

doc.add_heading('Crear un nuevo pedido:', level=4)
add_numbered('Haga clic en una mesa con estado "Libre".')
add_numbered('Se abrirá un panel lateral (drawer) para crear el pedido.')
add_numbered('Seleccione los platos del menú que desea agregar (cantidad y notas opcionales).')
add_numbered('Presione "Crear Pedido". La mesa cambiará automáticamente a estado "Ocupada".')

doc.add_heading('Agregar items a un pedido existente:', level=4)
add_numbered('Haga clic en una mesa "Ocupada".')
add_numbered('En el panel lateral, use el botón "Agregar Item" para añadir más platos.')
add_numbered('Seleccione los nuevos items y presione "Agregar al Pedido".')

add_note(
    'Puede agregar notas especiales a cada item (ej: "sin cebolla", "bien cocido").',
    'tip'
)

doc.add_heading('5.1.2 Kanban de Pedidos', level=3)
add_para('La página de Pedidos muestra un tablero Kanban con cuatro columnas '
         'que representan el flujo de trabajo de cada pedido:')

add_info_table(
    ['Columna', 'Descripción'],
    [
        ['Pendientes', 'Pedidos recién creados, esperando ser tomados por cocina'],
        ['En Preparación', 'Pedidos que están siendo cocinados'],
        ['Listos', 'Pedidos terminados, esperando ser entregados a la mesa'],
        ['Entregados', 'Pedidos ya servidos al cliente'],
    ]
)

add_para(
    'Puede arrastrar y soltar (drag & drop) las tarjetas de pedidos entre columnas '
    'para actualizar su estado. Como mesero, puede mover pedidos de "Listos" a '
    '"Entregados" cuando sirva los platos.'
)

add_bold_normal('Para entregar un pedido: ', 'Arrastre la tarjeta de la columna "Listos" a "Entregados".')
add_para('')
add_bold_normal('Para ver detalles: ', 'Haga clic en cualquier tarjeta para ver los items del pedido.')

page_break()

# ─── 5.2 COCINA ───
doc.add_heading('5.2 Cocina', level=2)

add_para(
    'El personal de cocina visualiza los pedidos entrantes y gestiona la preparación '
    'de cada plato de forma individual.',
    bold=True
)

doc.add_heading('5.2.1 Tablero de Cocina', level=3)
add_para('La página de Cocina muestra un Kanban similar al del mesero, pero centrado '
         'en los items (platos) individuales, no en el pedido completo:')

add_info_table(
    ['Columna', 'Descripción', 'Acción del Cocinero'],
    [
        ['Pendientes', 'Items nuevos esperando ser cocinados', 'Tomar el siguiente'],
        ['En Preparación', 'Items que se están cocinando', 'Marcar como listo al terminar'],
        ['Listos', 'Items terminados para ser entregados', 'Notificar al mesero'],
    ]
)

add_para('Flujo de trabajo en cocina:')

add_numbered('Los pedidos nuevos aparecen automáticamente en "Pendientes".')
add_numbered('Arrastre un item a "En Preparación" cuando comience a cocinarlo.')
add_numbered('Arrastre a "Listos" cuando el plato esté terminado.')
add_para('')
add_para(
    'Cuando todos los items de un pedido están en "Listos", el sistema notifica '
    'automáticamente a los meseros. También puede marcar items individuales como '
    '"Preparando" o "Listo" usando los botones en cada tarjeta.'
)

add_note(
    'Los items se pueden marcar como "Cancelados" si el mesero lo solicita '
    '(Ej: el cliente cambió su pedido).',
    'info'
)

page_break()

# ─── 5.3 CAJERO ───
doc.add_heading('5.3 Cajero', level=2)

add_para(
    'El cajero gestiona la facturación, el cobro y los reportes de ventas.',
    bold=True
)

doc.add_heading('5.3.1 Panel de Caja', level=3)
add_para('La página de Caja muestra un dashboard con información resumida:')
add_bullet('Total facturado en el día')
add_bullet('Total cobrado')
add_bullet('Cuentas pendientes de pago')
add_bullet('Pagos recientes')

doc.add_heading('5.3.2 Facturar un Pedido', level=3)
add_numbered('En la sección "Pedidos para Facturar", seleccione el pedido completado.')
add_numbered('Presione "Generar Cuenta". El sistema calcula automáticamente:')
add_bullet('Subtotal: suma de todos los items del pedido', level=1)
add_bullet('IVA (10%): impuesto sobre el subtotal', level=1)
add_bullet('Total: subtotal + IVA', level=1)
add_numbered('Una vez generada la cuenta, presione "Pagar" para registrar el cobro.')

doc.add_heading('5.3.3 Procesar un Pago', level=3)
add_para('Al presionar "Pagar", seleccione el método de pago:')

add_info_table(
    ['Método', 'Descripción'],
    [
        ['Efectivo', 'Registre el monto con el que paga el cliente. El sistema calcula el vuelto automáticamente.'],
        ['Tarjeta', 'Pago con débito o crédito. No requiere cálculo de vuelto.'],
        ['Transferencia', 'Pago por transferencia bancaria o Mercado Pago.'],
    ]
)

add_para('')
add_bold_normal('Paso a paso para pago en efectivo: ', '')
add_numbered('Seleccione "Efectivo" como método de pago.')
add_numbered('Ingrese el monto con el que paga el cliente.')
add_numbered('El sistema muestra el vuelto a entregar.')
add_numbered('Confirme el pago. La cuenta se marca como "Pagada".')

add_para('')
add_bold_normal('Paso a paso para pago con tarjeta o transferencia: ', '')
add_numbered('Seleccione "Tarjeta" o "Transferencia".')
add_numbered('Confirme el pago. No requiere cálculo de vuelto.')

add_para('')
add_para('Al confirmar el pago:')
add_bullet('La mesa pasa automáticamente a estado "Limpieza"')
add_bullet('Se genera la factura en PDF para imprimir o descargar')
add_bullet('El pedido queda registrado en el historial de ventas')

doc.add_heading('5.3.4 Factura en PDF', level=3)
add_para(
    'Cada factura incluye: logo de la pizzería, número de mesa, lista de items con '
    'cantidades y precios, subtotal, IVA (10%), total, método de pago, y un código '
    'QR que enlaza a la vista pública de la cuenta.'
)

doc.add_heading('5.3.5 Reporte de Ventas', level=3)
add_para(
    'El cajero puede generar reportes de ventas en PDF con la siguiente información:'
)
add_bullet('Resumen general: total de ventas, cantidad de transacciones, ticket promedio')
add_bullet('Desglose por método de pago (efectivo, tarjeta, transferencia)')
add_bullet('Historial detallado de todas las ventas del período')
add_bullet('Total general al final del reporte')
add_para('')
add_para('Puede filtrar por rango de fechas para obtener reportes diarios, semanales o mensuales.')

doc.add_heading('5.3.6 Cierre del Día', level=3)
add_para(
    'Al finalizar la jornada, el cajero puede realizar el cierre de caja. '
    'Esta acción marca todas las cuentas pagadas como "Cerradas" y prepara '
    'el sistema para un nuevo día de operaciones.'
)

add_note(
    'El cierre del día es una operación importante. Asegúrese de que todos los '
    'pagos estén registrados antes de realizarlo.',
    'caution'
)

page_break()

# ─── 5.4 ADMINISTRADOR ───
doc.add_heading('5.4 Administrador', level=2)

add_para(
    'El administrador tiene acceso completo a todas las funcionalidades del sistema, '
    'más la gestión del menú y del panel de administración de Django.',
    bold=True
)

doc.add_heading('5.4.1 Gestión del Menú', level=3)
add_para('La página de Menú permite administrar las categorías y los items del menú del restaurante.')

doc.add_heading('Categorías:', level=4)
add_bullet('Crear: Agregue nuevas categorías (Ej: "Pizzas", "Bebidas", "Postres")')
add_bullet('Editar: Modifique el nombre o el orden de visualización')
add_bullet('Eliminar: Borre categorías (solo si no tienen items asociados)')

doc.add_heading('Items del Menú:', level=4)
add_bullet('Crear: Agregue nuevos platos con nombre, descripción, precio, categoría y disponibilidad')
add_bullet('Editar: Modifique cualquier atributo de un plato existente')
add_bullet('Eliminar: Borre platos del menú')
add_bullet('Disponibilidad: Active o desactive platos temporalmente sin eliminarlos')

add_note(
    'Los items marcados como "No disponible" no aparecerán en el menú al crear pedidos.',
    'tip'
)

doc.add_heading('5.4.2 Panel de Administración Django', level=3)
add_para(
    'El administrador también puede acceder al panel de administración de Django '
    'en la ruta /admin/. Allí puede gestionar:'
)
add_bullet('Usuarios: Crear, modificar roles y permisos')
add_bullet('Menú: Gestión avanzada de categorías e items')
add_bullet('Mesas: Administración completa de mesas')
add_bullet('Pedidos: Visualización y edición de todos los pedidos')
add_bullet('Facturación: Historial completo de facturas')
add_bullet('Dispositivos: Ver tokens de dispositivos registrados para notificaciones')

page_break()

# ═══════════════════════════════════════════════════════════
# 6. NOTIFICACIONES EN TIEMPO REAL
# ═══════════════════════════════════════════════════════════

doc.add_heading('6. Notificaciones en Tiempo Real', level=1)

add_para(
    'El sistema cuenta con un sistema de notificaciones en tiempo real que mantiene '
    'a todo el personal informado de los cambios importantes sin necesidad de recargar '
    'la página.'
)

doc.add_heading('6.1 WebSockets', level=2)
add_para(
    'La aplicación utiliza WebSockets para recibir actualizaciones instantáneas. '
    'Cuando ocurre un evento relevante (nuevo pedido, item listo, etc.), la '
    'interfaz se actualiza automáticamente.'
)

add_info_table(
    ['Evento', '¿Quién lo genera?', '¿Quién lo recibe?'],
    [
        ['Nuevo pedido creado', 'Mesero', 'Cocina (vía WebSocket + Push)'],
        ['Item comenzó a prepararse', 'Cocina', 'Mesero del pedido'],
        ['Item terminado (listo)', 'Cocina', 'Mesero del pedido'],
        ['Pedido listo', 'Cocina (auto)', 'Todos los meseros'],
        ['Pedido entregado', 'Mesero', 'Cajero'],
        ['Cuenta generada', 'Cajero', 'Cajero (actualiza dashboard)'],
        ['Pago registrado', 'Cajero', 'Cajero (actualiza dashboard)'],
    ]
)

doc.add_heading('6.2 Notificaciones Push (Opcional)', level=2)
add_para(
    'En dispositivos móviles y en la web, el sistema puede enviar notificaciones '
    'push usando Firebase Cloud Messaging (FCM). Esto permite recibir alertas '
    'incluso cuando la aplicación no está en primer plano.'
)

add_note(
    'Para recibir notificaciones push en Android, la aplicación debe ser instalada '
    'desde un archivo APK o desde la Play Store.',
    'info'
)

page_break()

# ═══════════════════════════════════════════════════════════
# 7. MANUAL DEL BACKEND
# ═══════════════════════════════════════════════════════════

doc.add_heading('7. Manual del Backend (API)', level=1)

add_para(
    'Esta sección está dirigida a desarrolladores y administradores que necesiten '
    'interactuar directamente con la API REST del sistema.'
)

doc.add_heading('7.1 Autenticación', level=2)
add_para('La API utiliza JWT (JSON Web Tokens) para autenticación:')

add_code_block(
    '# Obtener token:\n'
    'POST /api/auth/login/\n'
    '{\n'
    '  "username": "mesero1",\n'
    '  "password": "contraseña"\n'
    '}\n\n'
    '# Respuesta:\n'
    '{\n'
    '  "access": "eyJ...",\n'
    '  "refresh": "eyJ..."\n'
    '}'
)

add_para('')
add_bold_normal('Access Token: ', 'Duración de 4 horas. Se envía en el header Authorization.')
add_para('')
add_bold_normal('Refresh Token: ', 'Duración de 1 día. Se usa para renovar el access token.')

add_code_block(
    '# Usar token en requests:\n'
    'Authorization: Bearer eyJ...\n\n'
    '# Refrescar token:\n'
    'POST /api/auth/refresh/\n'
    '{\n'
    '  "refresh": "eyJ..."\n'
    '}'
)

doc.add_heading('7.2 Endpoints Principales', level=2)

add_info_table(
    ['Endpoint', 'Método', 'Descripción'],
    [
        ['/api/auth/login/', 'POST', 'Iniciar sesión (obtener JWT)'],
        ['/api/auth/me/', 'GET', 'Obtener datos del usuario actual'],
        ['/api/menu/categories/', 'GET', 'Listar categorías del menú'],
        ['/api/menu/items/', 'GET', 'Listar items del menú'],
        ['/api/tables/', 'GET', 'Listar todas las mesas'],
        ['/api/orders/', 'GET', 'Listar pedidos'],
        ['/api/orders/', 'POST', 'Crear un nuevo pedido'],
        ['/api/orders/{id}/add_item/', 'POST', 'Agregar item a pedido'],
        ['/api/bills/generate/', 'POST', 'Generar cuenta desde pedido'],
        ['/api/bills/{id}/pay/', 'POST', 'Registrar pago'],
        ['/api/bills/report/', 'GET', 'Obtener reporte de ventas'],
        ['/api/bills/close_day/', 'POST', 'Cerrar el día'],
    ]
)

doc.add_heading('7.3 Panel de Administración Django', level=2)
add_para(
    'El panel de administración está disponible en la ruta /admin/ del backend. '
    'Para acceder, inicie sesión con las credenciales de superusuario creadas '
    'durante la instalación.'
)

add_para('Desde el panel admin puede:')
add_bullet('Gestionar usuarios: crear, editar roles y permisos')
add_bullet('Administrar el menú completo (categorías e items)')
add_bullet('Ver y filtrar todas las mesas')
add_bullet('Inspeccionar pedidos y su estado')
add_bullet('Revisar el historial completo de facturación')
add_bullet('Ver dispositivos registrados para notificaciones push')

page_break()

# ═══════════════════════════════════════════════════════════
# 8. SOLUCIÓN DE PROBLEMAS
# ═══════════════════════════════════════════════════════════

doc.add_heading('8. Solución de Problemas', level=1)

doc.add_heading('8.1 Problemas Comunes', level=2)

add_info_table(
    ['Problema', 'Posible Causa', 'Solución'],
    [
        ['No puedo iniciar sesión', 'Credenciales incorrectas', 'Verifique usuario y contraseña. Contacte al administrador.'],
        ['La página no carga', 'Servidor apagado o error de conexión', 'Verifique que el backend esté corriendo. Revise la URL.'],
        ['Los pedidos no aparecen', 'Conexión WebSocket perdida', 'Recargue la página. Verifique su conexión a internet.'],
        ['No recibo notificaciones', 'Permisos no concedidos o FCM no configurado', 'Conceda permisos de notificación en el navegador.'],
        ['Error 500 en el servidor', 'Error interno del backend', 'Contacte al administrador con el mensaje de error exacto.'],
        ['El PDF no se descarga', 'Bloqueador de pop-ups activo', 'Desactive el bloqueador de pop-ups para este sitio.'],
        ['La app móvil no conecta', 'URL del API incorrecta', 'Verifique VITE_API_URL en la configuración de la app.'],
    ]
)

doc.add_heading('8.2 Logs del Sistema', level=2)
add_para('Para diagnosticar problemas técnicos:')

add_bold_normal('Backend: ', 'Los logs se muestran en la consola donde se ejecuta Django.')
add_code_block(
    '# Ver logs en tiempo real\n'
    'docker-compose logs -f backend\n\n'
    '# O en desarrollo local:\n'
    '# Los mensajes aparecen en la terminal donde corre runserver'
)

add_bold_normal('Frontend: ', 'Puede ver errores en la consola del navegador (F12 → Console).')

page_break()

# ═══════════════════════════════════════════════════════════
# 9. PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════════

doc.add_heading('9. Preguntas Frecuentes', level=1)

faqs = [
    (
        '¿Qué hago si un cliente quiere dividir la cuenta?',
        'Actualmente el sistema genera una cuenta única por pedido/mesa. '
        'Para dividir la cuenta, deberá crear pedidos separados por cliente '
        'o consultar con el administrador sobre opciones alternativas.'
    ),
    (
        '¿Puedo modificar un pedido después de crearlo?',
        'Sí, puede agregar o cancelar items individuales desde el panel de detalles '
        'del pedido. Los items ya en preparación deberán coordinarse con cocina.'
    ),
    (
        '¿Cómo agrego un nuevo usuario al sistema?',
        'Solo el administrador puede crear usuarios. Debe acceder al panel de '
        'administración de Django en /admin/ y crear un nuevo usuario asignándole un rol.'
    ),
    (
        '¿Qué pasa si se corta internet?',
        'El sistema requiere conexión a internet para funcionar. Si hay un corte, '
        'los datos no se perderán ya que están almacenados en el servidor. '
        'Al restablecerse la conexión, puede continuar trabajando normalmente.'
    ),
    (
        '¿Puedo usar el sistema en una tablet o celular?',
        'Sí, el frontend es completamente responsive y se adapta a cualquier tamaño '
        'de pantalla. Además, está disponible como aplicación Android nativa.'
    ),
    (
        '¿Cómo se calcula el IVA?',
        'El IVA se calcula automáticamente como el 10% del subtotal (suma de todos '
        'los items del pedido). Este valor puede ser configurado por el administrador '
        'si es necesario.'
    ),
    (
        '¿Los datos se respaldan automáticamente?',
        'El respaldo de la base de datos depende de la configuración del servidor. '
        'Consulte con su administrador sobre la política de backups.'
    ),
    (
        '¿Puedo personalizar el logo o los colores?',
        'Sí, el logo y los colores están configurados en los archivos del frontend. '
        'Un desarrollador puede modificarlos fácilmente.'
    ),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'❓ {pregunta}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_SECONDARY

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(respuesta)
    run2.font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════
# 10. GLOSARIO
# ═══════════════════════════════════════════════════════════

doc.add_heading('10. Glosario de Términos', level=1)

add_info_table(
    ['Término', 'Definición'],
    [
        ['JWT', 'JSON Web Token. Método de autenticación basado en tokens cifrados.'],
        ['API REST', 'Interfaz de programación que permite la comunicación entre frontend y backend.'],
        ['SPA', 'Single Page Application. Aplicación web que carga una sola página y se actualiza dinámicamente.'],
        ['WebSocket', 'Protocolo de comunicación bidireccional en tiempo real.'],
        ['Kanban', 'Tablero visual de gestión de tareas con columnas de estado.'],
        ['FCM', 'Firebase Cloud Messaging. Servicio de notificaciones push de Google.'],
        ['POS', 'Point of Sale. Sistema de punto de venta.'],
        ['DRF', 'Django REST Framework. Librería para construir APIs con Django.'],
        ['Zustand', 'Librería de gestión de estado para React.'],
        ['Vite', 'Herramienta de build para frontend moderna y rápida.'],
        ['Capacitor', 'Framework para convertir aplicaciones web en apps nativas.'],
        ['IVA', 'Impuesto al Valor Agregado. Tasa impositiva del 10%.'],
    ]
)

page_break()

# ═══════════════════════════════════════════════════════════
# CONTACTO Y SOPORTE
# ═══════════════════════════════════════════════════════════

doc.add_heading('Contacto y Soporte Técnico', level=1)

add_para(
    'Para consultas técnicas, reporte de errores o solicitudes de mejora, '
    'contacte al administrador del sistema o al equipo de desarrollo.'
)

add_para('')
add_info_table(
    ['Medio', 'Detalle'],
    [
        ['Administrador del Sistema', 'Consultar con el encargado del local'],
        ['Reporte de Errores', 'Informar al administrador con captura de pantalla y descripción del problema'],
        ['Solicitudes de Mejora', 'Contactar al equipo de desarrollo a través del administrador'],
    ]
)

add_para('')
add_para('')
add_para('')

# ─── Pie de página ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del Manual —')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = COLOR_PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Documento generado el {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(9)
run.font.color.rgb = COLOR_GRAY
run.italic = True


# ═══════════════════════════════════════════════════════════
# GUARDAR DOCUMENTO
# ═══════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), 'Manual de Usuario - Pizzería El Patio.docx')
doc.save(output_path)
print(f'[OK] Documento generado exitosamente:')
print(f'     {output_path}')
print(f'     Tamano: {os.path.getsize(output_path) / 1024:.1f} KB')
